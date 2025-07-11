from docx import Document
import PyPDF2
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import io
from fpdf import FPDF

def extract_text_from_file(file_storage):
    filename = file_storage.filename.lower()
    data = file_storage.read()
    if filename.endswith('.txt'):
        return data.decode('utf-8')
    elif filename.endswith('.docx'):
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    elif filename.endswith('.pdf'):
        try:
            pdf = PyPDF2.PdfReader(io.BytesIO(data))
            return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
        except Exception:
            images = convert_from_bytes(data)
            text = ""
            for img in images:
                text += pytesseract.image_to_string(img)
            return text
    else:
        return ''

def export_as_docx_or_pdf(text, format):
    if format == 'docx':
        return export_as_docx(text)
    else:
        return export_as_pdf(text)

def export_as_docx(text):
    from docx import Document
    import io
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    fake_file = io.BytesIO()
    doc.save(fake_file)
    fake_file.seek(0)
    return fake_file

def export_as_pdf(text):
    import io
    from fpdf import FPDF

    # Replace smart quotes and other problematic characters
    def clean_text(s):
        # Replace curly quotes/apostrophes with straight ones
        s = s.replace('\u2018', "'").replace('\u2019', "'")
        s = s.replace('\u201c', '"').replace('\u201d', '"')
        # Replace en dash and em dash with hyphen
        s = s.replace('\u2013', '-').replace('\u2014', '-')
        # Remove/replace any other non-latin1 chars with '?'
        s = ''.join((c if ord(c) < 256 else '?') for c in s)
        return s

    cleaned_text = clean_text(text)

    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 12)
            self.cell(0, 10, 'Anonymised Report', ln=True, align='C')
            self.ln(10)

    pdf = PDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for line in cleaned_text.split('\n'):
        pdf.multi_cell(0, 10, line)
    fake_file = io.BytesIO()
    pdf_bytes = pdf.output(dest='S').encode('latin1')
    fake_file.write(pdf_bytes)
    fake_file.seek(0)
    return fake_file

def export_as_txt(text):
    import io
    fake_file = io.BytesIO()
    fake_file.write(text.encode('utf-8'))
    fake_file.seek(0)
    return fake_file
