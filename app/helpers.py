import spacy
from faker import Faker
import re
import openai
import json
import random
from PyPDF2 import PdfReader
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_JUSTIFY
from collections import defaultdict

# --- Load SpaCy Model and Add Custom EntityRuler ---
try:
    nlp = spacy.load("en_core_web_lg")
except OSError:
    print("Downloading 'en_core_web_lg' model...")
    spacy.cli.download("en_core_web_lg")
    nlp = spacy.load("en_core_web_lg")

# Define more robust patterns
age_patterns = [
    # "45 years old", "23-year-old"
    {"label": "AGE", "pattern": [{"IS_DIGIT": True}, {"LOWER": {"IN": ["year", "years", "-year"]}, "OP": "+"},
                                 {"LOWER": "old", "OP": "?"}]},
    # "age 52", "aged fifty-one"
    {"label": "AGE", "pattern": [{"LOWER": {"IN": ["age", "aged"]}}, {"LIKE_NUM": True}]},
    # "fifty-two years old"
    {"label": "AGE", "pattern": [{"LIKE_NUM": True}, {"LOWER": {"IN": ["year", "years"]}, "OP": "+"}, {"LOWER": "old"}]}
]

money_patterns = [
    {"label": "MONEY", "pattern": [{"IS_CURRENCY": True}, {"LIKE_NUM": True}]},
    {"label": "MONEY", "pattern": [{"LIKE_NUM": True}, {"LOWER": {"IN": ["pounds", "dollars", "euros", "gbp", "usd"]}}]}
]

# Add EntityRuler to the pipeline before the standard "ner" component
ruler = nlp.add_pipe("entity_ruler", before="ner")
ruler.add_patterns(age_patterns)
ruler.add_patterns(money_patterns)

fake = Faker()


# --- File Extraction (no changes needed) ---
def extract_text_from_file(file):
    filename = file.filename
    if filename.endswith('.pdf'):
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif filename.endswith('.docx'):
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif filename.endswith('.txt'):
        return file.read().decode('utf-8')
    else:
        raise ValueError("Unsupported file type. Please upload a .txt, .pdf, or .docx file.")


# --- NEW: Smart Entity Grouping Function ---
def group_person_entities(ents):
    """
    Groups partial names under full names.
    e.g., "Sarah" and "Malik" will be linked to "Sarah Malik".
    Returns a dictionary mapping the canonical name to all its variations.
    """
    persons = sorted([ent for ent in ents if ent.label_ == 'PERSON'], key=lambda x: len(x.text), reverse=True)
    other_ents = [ent for ent in ents if ent.label_ != 'PERSON']

    grouped_persons = {}
    processed_texts = set()

    for p1 in persons:
        if p1.text in processed_texts:
            continue

        canonical_name = p1.text
        variations = {p1.text}
        processed_texts.add(p1.text)

        # Split full name into parts for matching
        name_parts = set(canonical_name.lower().split())

        for p2 in persons:
            if p2.text in processed_texts:
                continue

            # Check if p2 is a substring of p1 (e.g., "Sarah" in "Sarah Malik")
            # or if p2's parts are a subset of p1's parts.
            p2_parts = set(p2.text.lower().split())
            if p2_parts.issubset(name_parts):
                variations.add(p2.text)
                processed_texts.add(p2.text)

        grouped_persons[canonical_name] = list(variations)

    return grouped_persons, other_ents


# --- SpaCy Anonymization Logic (Upgraded) ---
def process_text_spacy(text, level):
    doc = nlp(text)

    # Corrected level map: AGE is a direct identifier, moved to 'low'
    level_map = {
        'low': {'PERSON', 'AGE'},
        'medium': {'PERSON', 'AGE', 'GPE', 'LOC', 'ORG', 'FAC'},
        'high': {'PERSON', 'AGE', 'GPE', 'LOC', 'ORG', 'FAC', 'DATE', 'TIME', 'MONEY'}
    }
    target_labels = level_map.get(level, set())

    # Filter entities based on the selected level
    relevant_ents = [ent for ent in doc.ents if ent.label_ in target_labels]

    # Group PERSON entities intelligently
    grouped_persons, other_ents = group_person_entities(relevant_ents)

    entities_to_review = []

    # Add grouped persons to the review list
    for canonical_name, variations in grouped_persons.items():
        # The 'text' to be replaced is a list of all variations
        # The 'display_text' is what the user sees in the stepper
        suggestions = get_entity_suggestions('PERSON', canonical_name)
        entities_to_review.append({
            'text_to_replace': variations,
            'display_text': canonical_name,
            'label': 'PERSON',
            'suggestions': suggestions
        })

    # Add other entities to the review list, ensuring no duplicates
    processed_other_texts = set()
    for ent in other_ents:
        if ent.text not in processed_other_texts:
            suggestions = get_entity_suggestions(ent.label_, ent.text)
            entities_to_review.append({
                'text_to_replace': [ent.text],
                'display_text': ent.text,
                'label': ent.label_,
                'suggestions': suggestions
            })
            processed_other_texts.add(ent.text)

    return entities_to_review


def finalize_anonymization_text(original_text, user_choices):
    anonymized_text = original_text
    clean_text = original_text

    # Sort choices to replace longer strings first, preventing substring issues
    sorted_choices = sorted(user_choices, key=lambda x: len(x['original']), reverse=True)

    for choice in sorted_choices:
        replacement = choice['replacement']

        # 'original_list' now contains all variations of an entity
        original_list = sorted(choice['original_list'], key=len, reverse=True)

        if replacement and original_list[0] != replacement:
            for original_variation in original_list:
                # Use regex word boundary to avoid partial matches
                anonymized_text = re.sub(r'\b' + re.escape(original_variation) + r'\b', f'<mark>{replacement}</mark>',
                                         anonymized_text)
                clean_text = re.sub(r'\b' + re.escape(original_variation) + r'\b', replacement, clean_text)

    return {
        'anonymized_text_highlighted': anonymized_text,
        'anonymized_text_clean': clean_text
    }


# --- Suggestion Generator (Corrected Age Suggestions) ---
def get_entity_suggestions(label, original_text=""):
    """Generates realistic suggestions for a given SpaCy entity label."""
    if label == 'PERSON':
        return [fake.name() for _ in range(3)]
    if label in ['GPE', 'LOC']:
        return [fake.city(), fake.country(), fake.address().split('\n')[0]]
    if label == 'ORG':
        return [fake.company() for _ in range(3)]
    if label == 'DATE':
        return [fake.date_this_decade().strftime('%B %d, %Y') for _ in range(3)]
    if label == 'FAC':
        facility_types = ['Centre', 'Building', 'Stadium', 'Clinic', 'Bridge', 'Plaza']
        return [f"{fake.street_name()} {random.choice(facility_types)}" for _ in range(3)]
    if label == 'MONEY':
        return [f"£{random.randint(10, 500)}", f"${random.randint(10, 500)}", f"€{random.randint(10, 500)}"]
    if label == 'TIME':
        return [fake.time(pattern="%I:%M %p") for _ in range(3)]
    if label == 'AGE':  # <-- CORRECTED LOGIC FOR AGE
        try:
            # Extract number from text like "fifty-two" or "52"
            num_part = [token.text for token in nlp(original_text) if token.like_num][0]
            # Convert word to number if needed
            if num_part.isalpha():
                # Note: this requires a library like 'word2number' for full support,
                # but for simplicity, we'll suggest random numbers as a robust fallback.
                return [str(random.randint(20, 70)) for _ in range(3)]
            else:
                base_age = int(num_part)
                return [str(base_age + random.randint(-5, 5)) for _ in range(3) if base_age > 5]
        except:
            return [str(random.randint(20, 70)) for _ in range(3)]

    return [f"[{label.lower()}_replacement]" for _ in range(3)]


# --- ChatGPT & File Generation Logic (no changes needed) ---
def anonymize_text_with_gpt(text, level, api_key):
    # This function remains the same as before
    client = openai.OpenAI(api_key=api_key)
    level_instructions = {
        'low': "Anonymize only personal names (PERSON) and ages (AGE).",
        'medium': "Anonymize personal names (PERSON), ages (AGE), geopolitical entities (GPE), locations (LOC), organizations (ORG), and facilities (FAC).",
        'high': "Anonymize all personal data including names, ages, locations, organizations, dates (DATE), times (TIME), and monetary values (MONEY)."
    }
    prompt = f"""
    You are an expert in text anonymization. Your task is to rewrite the following report to protect identity while preserving meaning.

    Instructions:
    1.  Read the "Original Report" below.
    2.  Anonymize it based on the specified "Anonymization Level": {level_instructions.get(level, "Comprehensive anonymization.")}
    3.  Correct any grammatical errors.
    4.  Generate realistic replacements.
    5.  Your final output must be a single JSON object with two keys: "anonymized_text" (the full rewritten report as a string) and "changes" (a list of objects, each with "original" and "replacement" keys). Do not include any text outside this JSON object.

    Original Report:
    ---
    {text}
    ---

    Provide the JSON output.
    """
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            response_format={"type": "json_object"}
        )
        response_data = json.loads(response.choices[0].message.content)
        anonymized_text = response_data.get('anonymized_text', '')
        changes = response_data.get('changes', [])
        highlighted_text = anonymized_text
        sorted_changes = sorted(changes, key=lambda x: len(x.get('replacement', '')), reverse=True)
        for change in sorted_changes:
            replacement = change.get('replacement')
            if replacement:
                highlighted_text = re.sub(r'\b' + re.escape(replacement) + r'\b', f'<mark>{replacement}</mark>',
                                          highlighted_text, 1)
        return {
            'anonymized_text_highlighted': highlighted_text,
            'anonymized_text_clean': anonymized_text
        }
    except openai.APIConnectionError as e:
        raise ConnectionError(f"OpenAI API connection error: {e}")
    except openai.RateLimitError as e:
        raise ConnectionError(f"OpenAI API rate limit exceeded: {e}")
    except openai.AuthenticationError as e:
        raise ValueError(f"OpenAI authentication failed. Check your API key. Error: {e}")
    except Exception as e:
        raise RuntimeError(f"An unexpected error occurred with the AI model: {e}")

def generate_pdf_report(text, is_highlighted):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    if 'BodyText' not in styles:
        styles.add(ParagraphStyle(
            name='BodyText',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=11,
            leading=14,
            alignment=TA_JUSTIFY,
        ))

    if is_highlighted:
        # --- FIX: Directly use the hex color string. ---
        # This resolves the AttributeError.
        text = text.replace('<mark>', '<font backColor="#FDFFB6">').replace('</mark>', '</font>')

    # The <br/> tag is the standard way to handle newlines in ReportLab Paragraphs.
    story = [Paragraph(line, styles['BodyText']) for line in text.replace('\n', '<br/>').split('<br/>')]
    doc.build(story)

    buffer.seek(0)
    return buffer

def generate_docx_report(text, is_highlighted):
    # This function remains the same as before
    document = Document()
    document.add_heading('Anonymized Report', level=1)
    if is_highlighted:
        segments = re.split(r'(<mark>.*?</mark>)', text)
        p = document.add_paragraph()
        for segment in segments:
            if segment.startswith('<mark>'):
                content = segment.replace('<mark>', '').replace('</mark>', '')
                run = p.add_run(content)
                run.font.highlight_color = 7  # Yellow in WD_COLOR_INDEX
            else:
                p.add_run(segment)
    else:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer